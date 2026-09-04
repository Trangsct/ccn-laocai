"""Soát dự thảo văn bản hành chính (.docx) - phần việc CƠ HỌC của dây chuyền.

Bạn chốt 04/9/2026: máy làm phần cơ học, Claude bản cao nhất đọc hiểu và chốt nội dung. Script này:

  1. Rút mọi văn bản được viện dẫn trong dự thảo (số, ký hiệu, ngày).
  2. Đối chiếu sổ đăng ký văn bản pháp luật của bộ plugin skill-sct: văn bản nào đã bị sửa đổi, thay thế,
     hoặc có hiệu lực ở tương lai gần - những chỗ dễ làm hỏng lập luận nhất.
  3. Soi mâu thuẫn NGAY TRONG tệp: cùng một văn bản ghi hai ngày khác nhau; câu ở Tờ trình và câu tương ứng
     ở dự thảo Quyết định nói khác nhau.
  4. Kiểm thể thức cơ bản theo Nghị định 30/2020/NĐ-CP: cỡ chữ, phông chữ, dấu "./." kết thúc, "Nơi nhận".
  5. Xuất `bao-cao-soat.md` (cho người đọc) và `trich-dan.json` (cho bước tải văn bản kèm theo).

Chạy:  python scripts/soat_du_thao.py du-thao/<tệp>.docx [--ra <thư mục>]
"""
import argparse
import csv
import difflib
import io
import json
import re
import sys
import unicodedata
import urllib.request
from datetime import date, datetime
from pathlib import Path

REGISTRY = ("https://raw.githubusercontent.com/Trangsct/skill-sct/main/registry/trang-thai.csv")
CHO = 60

# Các loại văn bản hay được viện dẫn; nhóm 1 = tên loại, 2 = số, 3 = ký hiệu, 4 = ngày
MAU_VIEN_DAN = re.compile(
    r"(Luật|Nghị định|Nghị quyết|Quyết định|Thông tư|Tờ trình|Báo cáo|Công văn|Kế hoạch|Chỉ thị)"
    r"\s+(?:số\s+)?(\d+[A-Za-zĐđ]?)\s*/\s*([0-9]{4}\s*/\s*[A-ZĐ\-]+|[A-ZĐ][A-ZĐ\-]*)"
    r"(?:[^.;\n]{0,40}?ngày\s+(\d{1,2}\s*/\s*\d{1,2}\s*/\s*\d{4}|\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}))?",
    re.IGNORECASE)


def bo_dau(s):
    s = unicodedata.normalize("NFD", str(s or ""))
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", s.replace("đ", "d").replace("Đ", "D")).strip().lower()


def doc_ngay(chu):
    """'15/3/2024' hoặc '15 tháng 3 năm 2024' -> date."""
    if not chu:
        return None
    chu = re.sub(r"\s+", "", chu)
    m = re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})", chu)
    if not m:
        m = re.match(r"(\d{1,2})tháng(\d{1,2})năm(\d{4})", chu)
    if not m:
        return None
    try:
        return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
    except ValueError:
        return None


def doc_docx(duong_dan):
    from docx import Document

    d = Document(str(duong_dan))
    doan = [p.text.strip() for p in d.paragraphs]
    for b in d.tables:
        for h in b.rows:
            for o in h.cells:
                doan += [p.text.strip() for p in o.paragraphs]
    return d, [t for t in doan if t]


def rut_vien_dan(doan):
    """Trả danh sách văn bản được viện dẫn, mỗi văn bản kèm mọi cách ghi ngày bắt gặp."""
    ra = {}
    for i, t in enumerate(doan):
        for m in MAU_VIEN_DAN.finditer(t):
            loai, so, ky_hieu, ngay = (x.strip() if x else x for x in m.groups())
            ky_hieu = re.sub(r"\s+", "", ky_hieu).upper()
            khoa = f"{so}/{ky_hieu}"
            muc = ra.setdefault(khoa, {"loai": loai.capitalize(), "so_ky_hieu": khoa,
                                       "ngay": [], "doan": [], "nguyen_van": []})
            if ngay and ngay not in muc["ngay"]:
                muc["ngay"].append(ngay)
            if i not in muc["doan"]:
                muc["doan"].append(i)
            muc["nguyen_van"].append(m.group(0)[:120])
    return list(ra.values())


def tai_so_dang_ky():
    """Sổ trạng thái văn bản của bộ plugin: {mã: dòng}. Không tải được thì bỏ qua bước đối chiếu."""
    try:
        req = urllib.request.Request(REGISTRY, headers={"User-Agent": "soat-du-thao"})
        with urllib.request.urlopen(req, timeout=CHO) as r:
            chu = r.read().decode("utf-8")
    except Exception as e:
        print(f"  (không tải được sổ đăng ký văn bản: {e})")
        return {}
    ra = {}
    for dong in csv.DictReader(io.StringIO(chu)):
        ma = (dong.get("ma") or "").strip()
        m = re.search(r"(\d+[A-Za-z]?)\s*/\s*(\d{4})", ma)
        if m:
            ra[f"{m.group(1)}/{m.group(2)}"] = dong          # ví dụ "32/2024"
    return ra


def doi_chieu_hieu_luc(vien_dan, so_dang_ky, hom_nay):
    canh_bao = []
    for vb in vien_dan:
        m = re.match(r"(\d+[A-Za-z]?)/(\d{4})", vb["so_ky_hieu"])
        dong = so_dang_ky.get(f"{m.group(1)}/{m.group(2)}") if m else None
        if not dong:
            continue
        ten = f"{vb['loai']} số {vb['so_ky_hieu']}"
        if (dong.get("bi_thay_the_boi") or "").strip():
            canh_bao.append(f"**{ten}** đã bị THAY THẾ bởi {dong['bi_thay_the_boi'].strip()} - "
                            f"xem lại toàn bộ chỗ viện dẫn.")
        if (dong.get("bi_sua_doi_boi") or "").strip():
            canh_bao.append(f"**{ten}** đã bị SỬA ĐỔI, BỔ SUNG bởi {dong['bi_sua_doi_boi'].strip()} - "
                            f"viện dẫn chuẩn phải ghi kèm văn bản sửa đổi, và kiểm tra điều khoản trích dẫn "
                            f"còn nguyên văn hay đã đổi.")
        hl = doc_ngay(dong.get("hieu_luc"))
        if hl and hl > hom_nay:
            canh_bao.append(f"**{ten}** có hiệu lực từ {dong['hieu_luc']}, tức SAU hôm nay - chỉ nhắc với "
                            f"tính chất chuẩn bị, không dùng làm căn cứ xử lý.")
        if (dong.get("du_thao_thay_the") or "").strip():
            canh_bao.append(f"{ten}: đang có dự thảo thay thế ({dong['du_thao_thay_the'].strip()}).")
    return canh_bao


def soi_ngay_lech(vien_dan):
    ra = []
    for vb in vien_dan:
        ngay = {re.sub(r"\s+", "", n) for n in vb["ngay"]}
        chuan = {doc_ngay(n) for n in ngay}
        if len(chuan - {None}) > 1:
            ra.append(f"**{vb['loai']} số {vb['so_ky_hieu']}** ghi {len(chuan)} ngày khác nhau trong cùng "
                      f"tệp: {', '.join(sorted(ngay))}")
    return ra


BO_QUA = {"ong", "ba", "-", ",", ".", ";", ":", "va", "cac", "la", "cua", "tai", "theo", "so",
          "to", "vien", "thanh", "kiem", "truong", "pho", "hoi", "dong", "giup", "viec"}


def chuan_hoa_cau(t):
    """Đưa câu về dạng so sánh được: bỏ dấu, bỏ số thứ tự đầu dòng, quy ngày về một kiểu, bỏ dấu câu.

    Nhờ vậy 'ngày 12/6/2025' và 'ngày 12 tháng 6 năm 2025' được coi là một, và '2. Ông A - Giám đốc' cũng
    giống '- Ông A, Giám đốc'; chỉ còn lại khác biệt THẬT về nội dung.
    """
    t = bo_dau(t)
    t = re.sub(r"(\d{1,2})\s*thang\s*(\d{1,2})\s*nam\s*(\d{4})", r"\1/\2/\3", t)
    t = re.sub(r"\b0(\d)/", r"\1/", t)
    t = re.sub(r"/0(\d)/", r"/\1/", t)
    t = re.sub(r"^\s*[-+*]|^\s*\d+(\.\d+)*[.)]", "", t)
    return re.sub(r"[^a-z0-9/ ]", " ", re.sub(r"\s+", " ", t)).strip()


def _khac_nhau(x, y):
    """Những từ thực sự khác nhau giữa hai câu (bỏ dấu câu và từ nối)."""
    a, b = bo_dau(x).split(), bo_dau(y).split()
    khac = []
    for the, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b).get_opcodes():
        if the != "equal":
            khac += a[i1:i2] + b[j1:j2]
    return [t for t in khac if t.strip(",.;:-") and t not in BO_QUA]


def soi_mau_thuan_noi_bo(doan, nguong_giong=0.80):
    """Câu ở Tờ trình và câu tương ứng ở dự thảo Quyết định gần giống nhau nhưng KHÁC chi tiết.

    Chỉ báo khi khác nhau ở NỘI DUNG (từ khác hoặc con số khác), bỏ qua các cặp chỉ khác cách chấm câu -
    nếu không thì mỗi tên người trong danh sách Hội đồng lại thành một cảnh báo vô ích.
    """
    dai = [(i, t) for i, t in enumerate(doan) if len(t) > 60]
    ra = []
    for a in range(len(dai)):
        for b in range(a + 1, len(dai)):
            x, y = dai[a][1], dai[b][1]
            cx, cy = chuan_hoa_cau(x), chuan_hoa_cau(y)
            if cx == cy:                 # chỉ khác cách trình bày -> không phải mâu thuẫn
                continue
            ty = difflib.SequenceMatcher(None, cx, cy).ratio()
            if not (nguong_giong <= ty < 0.995):
                continue
            khac = _khac_nhau(cx, cy)
            so_khac = set(re.findall(r"\d+", cx)) ^ set(re.findall(r"\d+", cy))
            if len(khac) >= 3:
                ra.append((ty, dai[a][0], dai[b][0], x, y, khac[:12], sorted(so_khac)))
    ra.sort(key=lambda r: -r[0])
    return ra[:6]


def kiem_the_thuc(tai_lieu, doan):
    """Vài lỗi thể thức hay gặp theo Nghị định 30/2020/NĐ-CP."""
    loi = []
    co_dau = any(re.search(r"\./\.\s*$", t) for t in doan)
    if not co_dau:
        loi.append("Không thấy dấu kết thúc **./.** ở cuối phần nội dung.")
    if not any(bo_dau(t).startswith("noi nhan") for t in doan):
        loi.append("Không thấy mục **Nơi nhận**.")
    trong = [t for t in doan if re.search(r"số\s*:?\s*\.{2,}|/TTr-SCT\s*$|ngày\s+\.{2,}", t, re.I)]
    if trong:
        loi.append(f"Còn {len(trong)} chỗ để trống số hoặc ngày (bình thường với bản trình ký, chỉ nhắc).")
    co, phong = set(), set()
    for p in tai_lieu.paragraphs:
        for r in p.runs:
            if r.font.size:
                co.add(round(r.font.size.pt))
            if r.font.name:
                phong.add(r.font.name)
    if co and not co & {13, 14}:
        loi.append(f"Cỡ chữ đang dùng {sorted(co)} - thể thức yêu cầu 13 hoặc 14.")
    la = {p for p in phong if "times" not in p.lower()}
    if la:
        loi.append(f"Có phông chữ lạ ngoài Times New Roman: {', '.join(sorted(la))}.")
    return loi


def bao_cao(ten_tep, vien_dan, canh_bao, ngay_lech, mau_thuan, the_thuc, doan):
    d = [f"# Báo cáo soát dự thảo: {ten_tep}", "",
         f"Máy soát lúc {datetime.now():%d/%m/%Y %H:%M}. Đây là phần CƠ HỌC; phần đọc hiểu và chốt nội dung "
         f"do Claude làm ở bước sau.", "",
         f"- Số đoạn văn bản: {len(doan)}", f"- Số văn bản được viện dẫn: {len(vien_dan)}", ""]

    d += ["## 1. Văn bản được viện dẫn", "", "| Loại | Số ký hiệu | Ngày ghi trong dự thảo | Số lần nhắc |",
          "|---|---|---|---|"]
    for vb in sorted(vien_dan, key=lambda v: v["so_ky_hieu"]):
        d.append(f"| {vb['loai']} | {vb['so_ky_hieu']} | {', '.join(vb['ngay']) or '(không ghi ngày)'} | "
                 f"{len(vb['nguyen_van'])} |")
    d.append("")

    def muc(tieu_de, dong, khi_rong):
        d.append(f"## {tieu_de}")
        d.append("")
        d.extend([f"- {x}" for x in dong] if dong else [khi_rong])
        d.append("")

    muc("2. Cảnh báo hiệu lực (đối chiếu sổ đăng ký văn bản của bộ plugin)", canh_bao,
        "Không thấy văn bản nào đã bị sửa đổi, thay thế hoặc chưa tới hiệu lực.")
    muc("3. Ngày tháng ghi lệch nhau trong cùng tệp", ngay_lech, "Không thấy.")
    muc("4. Thể thức", the_thuc, "Không thấy lỗi thể thức cơ bản.")

    d += ["## 5. Câu gần giống nhau nhưng khác chi tiết", "",
          "Thường là một nội dung viết ở hai nơi (Tờ trình và dự thảo Quyết định) mà không khớp nhau.", ""]
    if mau_thuan:
        for ty, i, j, x, y, khac, so_khac in mau_thuan:
            d += [f"**Giống {ty:.0%}** - đoạn {i} và đoạn {j}"
                  + (f" · số chỉ có ở một bên: {', '.join(so_khac)}" if so_khac else "") + ":", "",
                  f"- (A) {x}", f"- (B) {y}",
                  f"- Khác ở các từ: {', '.join(khac)}" if khac else "", ""]
    else:
        d += ["Không thấy cặp nào đáng ngờ.", ""]

    d += ["## 6. Việc của Claude ở bước sau", "",
          "- Đối chiếu nội dung trích dẫn với bản gốc từng văn bản nêu tại mục 1.",
          "- Kiểm tra mốc hiệu lực và quy định chuyển tiếp có ảnh hưởng tới lập luận không.",
          "- Kiểm tra số liệu, tên người, chức danh, tên đơn vị theo dữ liệu của Sở.",
          "- Sửa và bôi đỏ từng chỗ, kèm bảng giải trình.", ""]
    return "\n".join(d)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("tep", help="tệp .docx dự thảo")
    ap.add_argument("--ra", default=None, help="thư mục ghi kết quả (mặc định: cạnh tệp dự thảo)")
    a = ap.parse_args()

    tep = Path(a.tep)
    if not tep.exists():
        sys.exit(f"Không thấy tệp {tep}")
    ra = Path(a.ra) if a.ra else tep.parent / tep.stem
    ra.mkdir(parents=True, exist_ok=True)

    tai_lieu, doan = doc_docx(tep)
    vien_dan = rut_vien_dan(doan)
    print(f"Đọc {len(doan)} đoạn, thấy {len(vien_dan)} văn bản được viện dẫn.")

    so_dang_ky = tai_so_dang_ky()
    canh_bao = doi_chieu_hieu_luc(vien_dan, so_dang_ky, date.today())
    ngay_lech = soi_ngay_lech(vien_dan)
    mau_thuan = soi_mau_thuan_noi_bo(doan)
    the_thuc = kiem_the_thuc(tai_lieu, doan)

    (ra / "trich-dan.json").write_text(
        json.dumps({"tep": tep.name, "vien_dan": vien_dan}, ensure_ascii=False, indent=2), encoding="utf-8")
    (ra / "bao-cao-soat.md").write_text(
        bao_cao(tep.name, vien_dan, canh_bao, ngay_lech, mau_thuan, the_thuc, doan), encoding="utf-8")
    print(f"Đã ghi {ra}/bao-cao-soat.md và trich-dan.json")
    print(f"  Cảnh báo hiệu lực: {len(canh_bao)} | ngày lệch: {len(ngay_lech)} | "
          f"thể thức: {len(the_thuc)} | cặp câu đáng ngờ: {len(mau_thuan)}")


if __name__ == "__main__":
    main()
